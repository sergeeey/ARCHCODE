"""
Embed figures into ARCHCODE manuscript DOCX for bioRxiv submission.
Inserts numbered figures with captions at appropriate locations.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import sys

PLOTS_DIR = Path("D:/ДНК/plots")
INPUT_DOCX = Path("C:/Users/serge/Desktop/DNC/ARCHCODE_Manuscript_v2.1.docx")
OUTPUT_DOCX = Path("C:/Users/serge/Desktop/DNC/ARCHCODE_Manuscript_v2.1_with_figures.docx")

# Figure definitions: (figure_number, filename, caption, anchor_text)
# anchor_text = unique string in a paragraph AFTER which the figure is inserted
FIGURES = [
    (
        1,
        "positional_signal_95kb.png",
        "Figure 1. HBB within-category positional signal analysis (95 kb window). "
        "Mann–Whitney U test comparing pathogenic vs benign SSIM distributions within "
        "each functional category. All 1,103 ClinVar HBB variants cluster within 2.1 kb "
        "(2.2% of window). LR ΔAUC = −0.001 (p = 1.0).",
        "occupy the same narrow genomic interval",
    ),
    (
        2,
        "positional_signal_cftr.png",
        "Figure 2. CFTR within-category positional signal analysis (317 kb window, 3,349 variants). "
        "Despite 63.6% TAD coverage and 29-fold greater positional diversity than HBB, "
        "within-category SSIM does not predict pathogenicity. LR ΔAUC = +0.008 (p = 1.0). "
        "SSIM range compressed to 0.9948–1.0000 due to matrix-size dilution at 317×317.",
        "CFTR ClinVar dataset",
    ),
    (
        3,
        "positional_signal_tp53.png",
        "Figure 3. TP53 within-category positional signal analysis (300 kb window, 2,795 variants). "
        "Variant spread 109.9 kb (36.6% of window). LR ΔAUC = +0.023 (p = 0.65) — smallest "
        "p-value among loci but firmly non-significant. Frameshift (n=534, all pathogenic) "
        "shows strongest disruption (mean SSIM = 0.9964).",
        "matrix-size dilution at 300",
    ),
    (
        4,
        "positional_signal_brca1.png",
        "Figure 4. BRCA1 within-category positional signal analysis (400 kb window, 10,682 variants). "
        "Largest ClinVar cohort. LR ΔAUC = −0.0002 (p = 1.0) — most decisive null result. "
        "Synonymous (n=5,520) MW-U p = 0.53, completely null despite massive statistical power. "
        "SSIM range 0.9982–1.0000.",
        "consistent with the dilution interpretation",
    ),
    (
        5,
        "tda_proof_of_concept_95kb.png",
        "Figure 5. TDA proof-of-concept: HBB locus (95 kb). Persistent homology (H1 loops) "
        "captures topological changes in ARCHCODE contact matrices under variant perturbation. "
        "Spearman ρ = −0.96 (p = 0.003) between SSIM and Wasserstein H1 distance.",
        "best captures the dominant regulatory architecture",
    ),
    (
        6,
        "tda_proof_of_concept_cftr.png",
        "Figure 6. TDA proof-of-concept: CFTR locus (317 kb). Perfect rank correlation "
        "(ρ = −1.00) between SSIM and Wasserstein H1 distance across variant categories.",
        None,  # inserted right after Figure 5
    ),
    (
        7,
        "tda_proof_of_concept_tp53.png",
        "Figure 7. TDA proof-of-concept: TP53 locus (300 kb). Weaker but significant "
        "correlation (ρ = −0.85, p = 0.015). 7 H1 persistent homology features in wild-type, "
        "reflecting greater structural complexity than HBB.",
        None,  # inserted right after Figure 6
    ),
    (
        8,
        "tda_proof_of_concept_brca1.png",
        "Figure 8. TDA proof-of-concept: BRCA1 locus (400 kb). TDA sensitivity drops at "
        "400×400 resolution: ρ = NaN (category-level), ρ = −0.21 (positional, p = 0.43, ns). "
        "8 H1 features in wild-type but perturbations too small to shift topology.",
        None,  # inserted right after Figure 7
    ),
    (
        9,
        "bayesian_fit_history.png",
        "Figure 9. Bayesian parameter optimization: trial history (Optuna GPSampler, 200 trials). "
        "Objective function: mean Pearson r between ARCHCODE wild-type and K562 Hi-C. "
        "Δr = +0.0001 — negligible improvement over grid-search parameters.",
        "not as a failed optimization",
    ),
    (
        10,
        "bayesian_fit_contours.png",
        "Figure 10. Bayesian optimization: parameter space contour plots. All three parameters "
        "(α, γ, k_base) converge to lower bounds, indicating the optimizer minimizes the "
        "Kramer kinetics term entirely. k_base importance = 90% (fANOVA).",
        None,  # right after Figure 9
    ),
    (
        11,
        "bayesian_fit_importance.png",
        "Figure 11. Bayesian optimization: parameter importance (fANOVA). k_base accounts for "
        "90% of objective variance; α (5%) and γ (5%) contribute negligibly. Hi-C correlation "
        "is architecture-driven, not kinetics-driven.",
        None,  # right after Figure 10
    ),
]


def find_paragraph_index(doc: Document, anchor: str) -> int | None:
    """Find the index of the paragraph containing anchor text."""
    for i, para in enumerate(doc.paragraphs):
        if anchor in para.text:
            return i
    return None


def insert_figure_after_paragraph(doc: Document, para_idx: int, img_path: str,
                                   caption: str, width_inches: float = 6.0):
    """Insert a figure with caption after the paragraph at para_idx.

    Returns the index of the last inserted element (caption paragraph).
    """
    # python-docx doesn't have a direct "insert after" — we manipulate the XML
    from docx.oxml.ns import qn
    import copy

    ref_para = doc.paragraphs[para_idx]
    ref_element = ref_para._element

    # Create image paragraph
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    img_element = img_para._element

    # Create caption paragraph
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_element = cap_para._element

    # Also add a blank paragraph for spacing
    spacer = doc.add_paragraph()
    spacer_element = spacer._element

    # Move elements: they were appended at the end, now relocate after ref
    body = doc.element.body
    # Remove from end
    body.remove(img_element)
    body.remove(cap_element)
    body.remove(spacer_element)
    # Insert after reference paragraph
    ref_element.addnext(spacer_element)
    ref_element.addnext(cap_element)
    ref_element.addnext(img_element)


def main():
    doc = Document(str(INPUT_DOCX))
    print(f"Loaded: {INPUT_DOCX}")
    print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    # Track insertion offset (each insert adds 3 paragraphs)
    offset = 0
    last_insert_idx = None

    for fig_num, filename, caption, anchor in FIGURES:
        img_path = PLOTS_DIR / filename
        if not img_path.exists():
            print(f"  [SKIP] Figure {fig_num}: {filename} not found")
            continue

        if anchor is not None:
            idx = find_paragraph_index(doc, anchor)
            if idx is None:
                print(f"  [WARN] Figure {fig_num}: anchor not found: '{anchor[:50]}...'")
                continue
            insert_idx = idx
        else:
            # Insert right after previous figure (use last_insert_idx + 3)
            if last_insert_idx is None:
                print(f"  [WARN] Figure {fig_num}: no anchor and no previous figure")
                continue
            insert_idx = last_insert_idx + 3  # img + caption + spacer

        insert_figure_after_paragraph(doc, insert_idx, str(img_path), caption)
        last_insert_idx = insert_idx
        print(f"  [OK] Figure {fig_num}: {filename} → after para {insert_idx}")

    doc.save(str(OUTPUT_DOCX))
    print(f"\nSaved: {OUTPUT_DOCX}")
    print(f"Final paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
